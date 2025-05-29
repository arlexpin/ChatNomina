import requests
import tempfile
from docx import Document
import logging
import os

logger = logging.getLogger(__name__)

def cargar_archivos_txt_desde_sharepoint(archivos_json):
    cache = {}
    for archivo in archivos_json:
        nombre = archivo.get("name", "").strip()
        if nombre.lower().endswith(".txt"):
            url = archivo.get("@microsoft.graph.downloadUrl")
            if not url:
                continue
            try:
                response = requests.get(url)
                if response.status_code == 200:
                    # Procesar el archivo como datos estructurados
                    lineas = []
                    for linea in response.content.decode("utf-8").splitlines():
                        if linea.strip():
                            # Dividir por punto y coma y limpiar cada campo
                            campos = [campo.strip() for campo in linea.split(";")]
                            lineas.append(campos)
                    if lineas:  # Solo agregar si hay contenido
                        cache[nombre] = lineas
                    else:
                        print(f"⚠️ Archivo vacío: {nombre}")
                else:
                    print(f"⚠️ Error descargando {nombre}: status {response.status_code}")
            except Exception as e:
                print(f"❌ Error procesando {nombre}: {e}")
    return cache

def cargar_documentos_word_desde_sharepoint(archivos_json):
    documentos = {}
    total_archivos = len([a for a in archivos_json if a.get("name", "").endswith((".doc", ".docx"))])
    logger.info(f"Encontrados {total_archivos} archivos Word para procesar")
    
    for archivo in archivos_json:
        nombre = archivo.get("name", "")
        if nombre.endswith(".doc") or nombre.endswith(".docx"):
            url = archivo.get("@microsoft.graph.downloadUrl")
            if not url:
                logger.warning(f"No se encontró URL de descarga para {nombre}")
                continue
                
            try:
                logger.info(f"Descargando {nombre}...")
                response = requests.get(url)
                if response.status_code != 200:
                    logger.error(f"Error descargando {nombre}: status {response.status_code}")
                    continue
                    
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    tmp.write(response.content)
                    tmp.flush()
                    logger.debug(f"Archivo temporal creado: {tmp.name}")
                    
                    try:
                        doc = Document(tmp.name)
                        texto = []
                        total_parrafos = len(doc.paragraphs)
                        total_tablas = len(doc.tables)
                        
                        logger.info(f"Procesando {nombre}: {total_parrafos} párrafos, {total_tablas} tablas")

                        # Párrafos normales
                        for i, p in enumerate(doc.paragraphs, 1):
                            if p.text.strip():
                                texto.append(p.text.strip())
                                if i % 100 == 0:
                                    logger.debug(f"Procesados {i}/{total_parrafos} párrafos en {nombre}")

                        # Contenido de tablas
                        for i, tabla in enumerate(doc.tables, 1):
                            for fila in tabla.rows:
                                fila_texto = " | ".join(cell.text.strip() for cell in fila.cells if cell.text.strip())
                                if fila_texto:
                                    texto.append(fila_texto)
                            logger.debug(f"Procesada tabla {i}/{total_tablas} en {nombre}")

                        contenido = "\n".join(texto)
                        if contenido.strip():
                            documentos[nombre] = contenido
                            logger.info(f"Documento {nombre} procesado exitosamente: {len(texto)} fragmentos")
                        else:
                            logger.warning(f"Documento {nombre} está vacío después del procesamiento")
                            
                    except Exception as e:
                        logger.error(f"Error procesando contenido de {nombre}: {str(e)}", exc_info=True)
                    finally:
                        try:
                            os.unlink(tmp.name)
                            logger.debug(f"Archivo temporal {tmp.name} eliminado")
                        except Exception as e:
                            logger.warning(f"Error eliminando archivo temporal {tmp.name}: {str(e)}")

            except Exception as e:
                logger.error(f"Error general procesando {nombre}: {str(e)}", exc_info=True)
                
    logger.info(f"Procesamiento de Word completado: {len(documentos)}/{total_archivos} documentos cargados")
    return documentos